import os
from pathlib import Path

import docx

from Practice6.Utils.pdf2txt import pdf_to_text


class ExtensionError(Exception):
    pass


class File:
    _READABLE_FORMATS: list = None

    def __init__(self, path):
        self.path: Path = Path(path)
        self.stringify_path: str = str(self.path)
        if self.path.exists() and not self.path.is_file():
            raise TypeError(f'{self.path.name} not a file')
        self.name, self.ext = os.path.splitext(self.stringify_path)
        if self.ext not in self._READABLE_FORMATS:
            raise ExtensionError(f'Invalid extension: {self.ext}. Required: {self._READABLE_FORMATS}')
        self.abs_path = self.path.absolute()

    def _validate_reading(self):
        if not self.path.exists():
            raise FileNotFoundError(f'Not found {self.abs_path}')

    @classmethod
    def get_readable_formats(cls):
        return cls._READABLE_FORMATS

    def read_lines_iter(self):
        pass

    def read_lines(self):
        pass

    def __str__(self):
        return str(self.abs_path)


class TextFile(File):
    _READABLE_FORMATS = ['.txt', '.c', '']

    def read_lines_iter(self):
        self._validate_reading()
        with open(self.stringify_path) as f:
            while True:
                line = f.readline()
                if line == '':
                    break
                yield line

    def read_lines(self):
        self._validate_reading()
        with open(self.stringify_path) as f:
            return f.readlines()


class PdfFile(File):
    _READABLE_FORMATS = ['.pdf']

    def __init__(self, path, decode: str = "utf-8"):
        super().__init__(path)
        self.decode: str = decode

    def read_lines_iter(self):
        for line in self.read_lines():
            yield line

    def read_lines(self):
        self._validate_reading()
        raw_text_data = pdf_to_text(self.stringify_path)
        decoded_text = raw_text_data.decode(self.decode)
        return decoded_text.splitlines()


class DocxFile(File):
    _READABLE_FORMATS = ['.docx']

    def read_lines_iter(self):
        self._validate_reading()
        doc_object = docx.Document(self.stringify_path)
        for paragraph in doc_object.paragraphs:
            if paragraph.text:
                yield paragraph.text

    def read_lines(self):
        self._validate_reading()
        doc_object = docx.Document(self.stringify_path)
        text = [paragraph.text for paragraph in doc_object.paragraphs if paragraph.text]
        return text


class DataLoader:
    __AVAILABLE_CLASSES__ = [TextFile, PdfFile, DocxFile]
    __AVAILABLE_EXTENSIONS__ = {}
    for _class in __AVAILABLE_CLASSES__:
        for extension in _class.get_readable_formats():
            __AVAILABLE_EXTENSIONS__[extension] = _class

    def __init__(self, path, formats):
        self.path: Path = Path(path)
        self.stringify_path = str(self.path)
        self.is_exists = self.path.exists()
        self.abs_path = self.path.absolute()
        self.stringify_abs_path = str(self.abs_path)

        difference_extension = set(formats) - set(self.__AVAILABLE_EXTENSIONS__.keys())
        if difference_extension:
            raise TypeError(f'Exception {difference_extension} is not supported')
        self.formats = formats
        if not self.is_exists:
            return
        self.is_file = self.path.is_file()
        self.is_dir = self.path.is_dir()
        if self.is_file:
            self.name, self.ext = os.path.splitext(self.stringify_path)
            self.reading_class = self._get_file_reading_class(self.ext)

    def _validate_reading_dir(self):
        if not self.is_dir:
            raise FileNotFoundError(f"{self.stringify_path} not a directory")

    def _get_file_reading_class(self, ext: str):
        try:
            return self.__AVAILABLE_EXTENSIONS__[ext]
        except KeyError:
            raise ExtensionError(
                f'Invalid extension: {ext}. Available extensions: {self.__AVAILABLE_EXTENSIONS__.keys()}')

    def __get_all_files_in_folders(self):
        self._validate_reading_dir()
        dir_name = self.path.name
        len_of_dir_name = len(dir_name)
        for root, subFolders, files in os.walk(self.stringify_path):
            depth = root[len_of_dir_name:].count(os.sep) + 1
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file_path)
                if ext not in self.formats:
                    continue
                reading_class = self._get_file_reading_class(ext)
                yield file_path, reading_class(file_path), depth

    def folder_walk_iter(self):
        for file_path, reader, depth in self.__get_all_files_in_folders():
            for line in reader.read_lines_iter():
                yield file_path, line, depth

    def folder_walk(self):
        for file_path, reader, depth in self.__get_all_files_in_folders():
            yield file_path, reader.read_lines(), depth

# if __name__ == '__main__':
#     d = DataLoader('wiki_articles')
#     for root, lines, depth in d.folder_walk():
#         print(root, depth)

# OUTPUT:
# wiki_articles\ml_frameworks\pytorch.txt 2
# wiki_articles\ml_frameworks\tensorflow.txt 2
# wiki_articles\programming_languages\compiled\c++.pdf 3
# wiki_articles\programming_languages\compiled\java.txt 3
# wiki_articles\programming_languages\interpreted\python.txt 3
# wiki_articles\programming_languages\interpreted\r.docx 3
