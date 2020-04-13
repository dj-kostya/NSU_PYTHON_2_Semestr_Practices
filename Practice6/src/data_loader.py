#!/usr/bin/env python

"""
 DataLoader pipeline. For reading data from different sources.

 Available sources:
    # from file
    # from all files in directory


 Features:
    TODO: add internet resources
"""
import os
from pathlib import Path

import docx

from Practice6 import pdf_to_text


class ExtensionError(Exception):
    """Checks extension errors. For example: not valid extension"""
    pass


class File:
    """
        Base class for working with file.

        Class has some useful methods for reading files.
    """
    _READABLE_FORMATS: list = None  # readable formats for this class.

    def __init__(self, path):
        self.path: Path = Path(path)
        self.stringify_path: str = str(self.path)
        if self.path.exists() and not self.path.is_file():
            raise TypeError(f'{self.path.name} not a file')
        self.name, self.ext = os.path.splitext(self.stringify_path)
        if self.ext not in self._READABLE_FORMATS:
            raise ExtensionError(f'Invalid extension: {self.ext}. Required: {self._READABLE_FORMATS}')
        self.abs_path = self.path.absolute()

    def _validate_reading(self):
        """Validate file for reading. Checking for existence of file."""
        if not self.path.exists():
            raise FileNotFoundError(f'Not found {self.abs_path}')

    @classmethod
    def get_readable_formats(cls):
        """Auxiliary method for getting readable formats of current class."""
        return cls._READABLE_FORMATS

    def read_lines_iter(self):
        """
        :return:
            Generator extension, that yields lines of file.
        """
        pass

    def read_lines(self):
        """
        :return:
            list of lines from the file.
        """
        pass

    def __str__(self):
        return str(self.abs_path)


class TextFile(File):
    """ Class for working with plain text files? that can readable with standard python functions. """

    _READABLE_FORMATS = ['.txt', '.c', '', '.py']

    def read_lines_iter(self):
        self._validate_reading()
        with open(self.stringify_path) as f:
            while True:
                line = f.readline()
                if line == '':
                    break
                yield line

    def read_lines(self):
        self._validate_reading()
        with open(self.stringify_path) as f:
            return f.readlines()


class PdfFile(File):
    """ Class for working with pdf files. Using pdf_to_text util. """
    _READABLE_FORMATS = ['.pdf']

    def __init__(self, path, decode: str = "utf-8"):
        super().__init__(path)
        self.decode: str = decode

    def read_lines_iter(self):
        for line in self.read_lines():
            yield line

    def read_lines(self):
        self._validate_reading()
        raw_text_data = pdf_to_text(self.stringify_path)
        decoded_text = raw_text_data.decode(self.decode)
        return decoded_text.splitlines()


class DocxFile(File):
    """ Class for working with docx files """
    _READABLE_FORMATS = ['.docx']

    def read_lines_iter(self):
        self._validate_reading()
        doc_object = docx.Document(self.stringify_path)
        for paragraph in doc_object.paragraphs:
            if paragraph.text:
                yield paragraph.text

    def read_lines(self):
        self._validate_reading()
        doc_object = docx.Document(self.stringify_path)
        text = [paragraph.text for paragraph in doc_object.paragraphs if paragraph.text]
        return text


class DataLoader:
    """
    Class for loading data from file or all files in a directory.
    """
    __AVAILABLE_CLASSES__ = [TextFile, PdfFile, DocxFile]
    __AVAILABLE_EXTENSIONS__ = {}
    for _class in __AVAILABLE_CLASSES__:
        for extension in _class.get_readable_formats():
            __AVAILABLE_EXTENSIONS__[extension] = _class

    def __init__(self, path, formats=None):
        self.path: Path = Path(path)
        self.stringify_path = str(self.path)
        self.is_exists = self.path.exists()
        self.abs_path = self.path.absolute()
        self.stringify_abs_path = str(self.abs_path)
        if not self.is_exists:
            return
        self.is_file = self.path.is_file()
        self.is_dir = self.path.is_dir()
        if self.is_file:
            self.name, self.ext = os.path.splitext(self.stringify_path)
            self.reading_class = self._get_file_reading_class(self.ext)(self.abs_path)
        if self.is_dir:
            difference_extension = set(formats) - set(self.__AVAILABLE_EXTENSIONS__.keys())
            if difference_extension:
                raise ExtensionError(f'Extensions: {difference_extension} are not supported')
            self.formats = formats

    def _validate_reading_dir(self):
        """ Function to validate a directory before reading"""
        if not self.is_dir:
            raise FileNotFoundError(f"{self.stringify_path} not a directory")

    def _get_file_reading_class(self, ext: str):
        """
        :param ext: string with extension of the file
        :return: Class for reading this file
        """
        try:
            return self.__AVAILABLE_EXTENSIONS__[ext]
        except KeyError:
            raise ExtensionError(
                f'Invalid extension: {ext}. Available extensions: {self.__AVAILABLE_EXTENSIONS__.keys()}')

    def __get_all_files_in_folders(self):
        """
        Auxiliary function to getting all files in a folder.
        :return:
        Generator expression, which yield file_path, class for reading this file, current recursion depth in file tree
        """
        self._validate_reading_dir()
        len_dir_name = len(self.stringify_path)
        for root, subFolders, files in os.walk(self.stringify_path):
            depth = root[len_dir_name:].count(os.sep) + 1
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file_path)
                if ext not in self.formats:
                    continue
                reading_class = self._get_file_reading_class(ext)
                yield file_path, reading_class(file_path), depth

    def folder_walk_iter(self):
        """
        Function for getting all files in a folder
        :return: Generator expression, which yield file_path of file, line of this file, current depth in file tree
        """
        for file_path, reader, depth in self.__get_all_files_in_folders():
            for line in reader.read_lines_iter():
                yield file_path, line, depth

    def folder_walk(self):
        """
        Function for getting all files in a folder
        :return: Generator expression, which yield file_path of file, all lines of that file, current depth in file tree
        """
        for file_path, reader, depth in self.__get_all_files_in_folders():
            yield file_path, reader.read_lines(), depth

    def read_lines_iter(self):
        """
        Function for reading all lines in file
        :return: Generator expression, which yield lines of this file
        """
        yield from self.reading_class.read_lines_iter()

    def read_lines(self):
        """
        Function for reading all lines in file
        :return: list of lines of this file
        """
        return self.reading_class.read_lines()

# region FOLDER_WALK TEST
########################################################################################################################
# CALL
# if __name__ == '__main__':
#     d = DataLoader('wiki_articles')
#     for root, lines, depth in d.folder_walk():
#         print(root, depth)
########################################################################################################################
# OUTPUT:
# wiki_articles\ml_frameworks\pytorch.txt 2
# wiki_articles\ml_frameworks\tensorflow.txt 2
# wiki_articles\programming_languages\compiled\c++.pdf 3
# wiki_articles\programming_languages\compiled\java.txt 3
# wiki_articles\programming_languages\interpreted\python.txt 3
# wiki_articles\programming_languages\interpreted\r.docx 3
########################################################################################################################
# endregion

# region ANOTHER TEST
########################################################################################################################
# CALL
# if __name__ == '__main__':
#     d = DataLoader('requirements.txt')
#     for line in d.read_lines_iter():
#         print(line, end='')
########################################################################################################################
# OUTPUT
# pdfminer.six
# python-docx
########################################################################################################################
# endregion
